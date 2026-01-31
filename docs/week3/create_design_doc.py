"""
CyberRisk Monitor Project Design Document Generator
Creates a comprehensive Word document meeting all rubric criteria
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Helper functions
def set_table_borders(table):
    """Set table borders - simplified approach"""
    pass  # Table Grid style handles borders

def add_image_with_caption(doc, image_path, caption, width=6):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f'[Image: {caption} - {image_path}]')

# ============================================
# TITLE PAGE
# ============================================
title = doc.add_heading('CyberRisk Monitor', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Lightweight Rule-Based Cybersecurity Monitoring and Risk Assessment Tool')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()

course = doc.add_paragraph('UMGC CMSC 495 - Computer Science Capstone')
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course.runs[0].font.size = Pt(12)

design_doc = doc.add_paragraph('Project Design Document')
design_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
design_doc.runs[0].font.size = Pt(14)
design_doc.runs[0].bold = True

semester = doc.add_paragraph('Spring 2026')
semester.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

team_label = doc.add_paragraph('Team Members:')
team_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_label.runs[0].bold = True

members = doc.add_paragraph('Mustafa Black-Castle, Daniel S. Garrett, Nicholas Porpora, and Cassandra Santacruz')
members.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================
# TABLE OF CONTENTS
# ============================================
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Introduction',
    '2. Project Scope',
    '3. Project Requirements',
    '4. Project Methodology',
    '5. Work Breakdown Structure and Project Tasks',
    '6. Project Schedule',
    '7. Project Resources',
    '8. Project Risks',
    '9. Project Evaluation Plan',
    '10. System Design',
    '11. Testing and Validation Strategy',
    '12. Conclusion',
    'Appendix A: Detailed Class Specifications',
    'Appendix B: Detection Rules Catalog',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ============================================
# 1. INTRODUCTION
# ============================================
doc.add_heading('1. Introduction', 1)

doc.add_paragraph(
    'This Project Design document presents the execution roadmap for CyberRisk Monitor, '
    'a lightweight rule-based cybersecurity monitoring and risk assessment tool developed '
    'as part of the UMGC CMSC 495 Computer Science Capstone. The purpose of this document '
    'is to translate the approved Project Plan into a structured design and implementation '
    'strategy that guides development, scheduling, risk management, and evaluation activities '
    'throughout the project lifecycle.'
)

doc.add_paragraph(
    'CyberRisk Monitor is designed as a proof-of-concept system rather than a full-scale '
    'enterprise solution. The project focuses on analyzing simulated or sample security logs '
    'using predefined detection rules to identify indicators of suspicious or potentially '
    'malicious activity. The system emphasizes clarity, interpretability, and accessibility, '
    'making it suitable for educational environments and small teams seeking an introductory '
    'cybersecurity monitoring capability.'
)

doc.add_paragraph(
    'This document builds upon the previously established Project Plan by expanding on '
    'technical requirements, task decomposition, scheduling logic, and evaluation criteria. '
    'It serves as a comprehensive reference for how the project will be executed, monitored, '
    'and assessed to ensure alignment with course objectives and successful capstone delivery.'
)

doc.add_heading('1.1 Document Purpose', 2)
doc.add_paragraph('This design document serves multiple purposes:')

purposes = [
    'Provide a detailed roadmap for project execution from Week 3 through Week 8',
    'Define clear scope boundaries, deliverables, and success criteria',
    'Establish a structured work breakdown with accurate time and resource estimates',
    'Identify and analyze potential risks with corresponding mitigation strategies',
    'Create measurable evaluation criteria aligned with project objectives',
    'Guide the development team through a systematic implementation process'
]
for purpose in purposes:
    doc.add_paragraph(purpose, style='List Bullet')

doc.add_heading('1.2 Intended Audience', 2)
doc.add_paragraph(
    'This document is intended for course instructors evaluating the project plan, '
    'team members executing the development work, and any stakeholders interested in '
    'understanding the project\'s scope and approach.'
)

# ============================================
# 2. PROJECT SCOPE
# ============================================
doc.add_page_break()
doc.add_heading('2. Project Scope', 1)

doc.add_paragraph(
    'The scope of the CyberRisk Monitor project defines the objectives, deliverables, '
    'boundaries, and limitations necessary to ensure successful completion within the '
    'fixed academic timeline of the capstone course.'
)

doc.add_heading('2.1 Project Objectives', 2)
doc.add_paragraph(
    'The primary objective is to design, implement, and deliver a functional cybersecurity '
    'monitoring tool capable of analyzing structured log data and identifying indicators of '
    'suspicious or potentially malicious activity using predefined detection rules. The following '
    'specific objectives guide the project:'
)

objectives_table = doc.add_table(rows=6, cols=3)
set_table_borders(objectives_table)
objectives_table.style = 'Table Grid'

objectives_table.cell(0, 0).text = 'Objective ID'
objectives_table.cell(0, 1).text = 'Objective Description'
objectives_table.cell(0, 2).text = 'Success Criteria'
for cell in objectives_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

objectives_data = [
    ('OBJ-01', 'Develop a log parsing system supporting multiple formats', 'Successfully parse syslog, JSON, and CSV formats with >95% accuracy'),
    ('OBJ-02', 'Implement rule-based detection engine', 'Execute 10-15 detection rules with correct pattern matching'),
    ('OBJ-03', 'Create risk scoring and categorization system', 'Accurately categorize events into Low/Medium/High/Critical levels'),
    ('OBJ-04', 'Build user-friendly reporting interface', 'Generate clear summary reports and visualizations'),
    ('OBJ-05', 'Deliver documented, maintainable codebase', 'Complete documentation with >80% code coverage in tests'),
]

for i, (obj_id, desc, criteria) in enumerate(objectives_data, 1):
    objectives_table.cell(i, 0).text = obj_id
    objectives_table.cell(i, 1).text = desc
    objectives_table.cell(i, 2).text = criteria

doc.add_paragraph()
cap = doc.add_paragraph('Table 2.1: Project Objectives and Success Criteria')
cap.runs[0].italic = True

doc.add_heading('2.2 Project Deliverables', 2)
doc.add_paragraph('The following deliverables will be produced upon project completion:')

deliverables = [
    'Functional CyberRisk Monitor application with all core modules',
    'LogParser module supporting syslog, JSON, and CSV formats',
    'RuleEngine module with 10-15 predefined detection rules',
    'RiskScorer module with configurable scoring algorithms',
    'ReportGenerator module producing summary reports and visualizations',
    'UserInterface providing command-line and/or graphical interaction',
    'Sample log datasets for testing and demonstration',
    'Comprehensive code documentation and user manual',
    'Final project presentation and demonstration materials'
]
for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('2.3 Scope Boundaries', 2)
doc.add_paragraph('The following table clearly defines what is included and excluded from this project:')

scope_table = doc.add_table(rows=8, cols=2)
set_table_borders(scope_table)
scope_table.style = 'Table Grid'

scope_table.cell(0, 0).text = 'In Scope'
scope_table.cell(0, 1).text = 'Out of Scope'
for cell in scope_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

in_scope = [
    'Analysis of simulated/sample security log data',
    'Implementation of 10-15 predefined detection rules',
    'Basic risk scoring and categorization (Low/Medium/High/Critical)',
    'Summary dashboards and structured reports',
    'Support for syslog, JSON, and CSV formats',
    'Command-line and/or simple graphical interface',
    'Offline batch processing of log files',
]

out_scope = [
    'Real-time monitoring of live network traffic',
    'Integration with production/enterprise environments',
    'Machine learning or AI-based detection techniques',
    'Automated incident response or remediation',
    'Enterprise-scale log aggregation or storage',
    'Regulatory compliance frameworks (HIPAA, PCI-DSS)',
    'Network packet capture or deep packet inspection',
]

for i, (in_s, out_s) in enumerate(zip(in_scope, out_scope), 1):
    scope_table.cell(i, 0).text = in_s
    scope_table.cell(i, 1).text = out_s

doc.add_paragraph()
cap = doc.add_paragraph('Table 2.2: Project Scope Boundaries')
cap.runs[0].italic = True

doc.add_heading('2.4 Constraints and Limitations', 2)

constraints_table = doc.add_table(rows=7, cols=2)
set_table_borders(constraints_table)
constraints_table.style = 'Table Grid'
constraints_table.cell(0, 0).text = 'Constraint Type'
constraints_table.cell(0, 1).text = 'Description'
for cell in constraints_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

constraints = [
    ('Time Constraint', 'Project must be completed within 6 weeks (Weeks 3-8 of the course)'),
    ('Team Size', 'Limited to 4 team members with varying availability'),
    ('Budget', 'Zero budget; all tools and resources must be free or open-source'),
    ('Technical', 'Cross-platform compatibility required (Windows, macOS, Linux)'),
    ('Data', 'Must use simulated/sample data only; no access to real security logs'),
    ('Complexity', 'Proof-of-concept level; not production-ready enterprise software'),
]

for i, (ctype, desc) in enumerate(constraints, 1):
    constraints_table.cell(i, 0).text = ctype
    constraints_table.cell(i, 1).text = desc

doc.add_paragraph()
cap = doc.add_paragraph('Table 2.3: Project Constraints and Limitations')
cap.runs[0].italic = True

doc.add_heading('2.5 Assumptions', 2)
assumptions = [
    'All team members have access to suitable development environments',
    'Team members possess foundational Python programming skills',
    'Adequate sample log data is available or can be generated for testing',
    'Team members can dedicate approximately 8-12 hours per week to the project',
    'Git/GitHub is accessible for version control and collaboration',
    'Internet access is available for research and tool downloads'
]
for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

# ============================================
# 3. PROJECT REQUIREMENTS
# ============================================
doc.add_page_break()
doc.add_heading('3. Project Requirements', 1)

doc.add_paragraph(
    'Project requirements for CyberRisk Monitor were derived from capstone course objectives, '
    'foundational cybersecurity practices, and the intended educational use of the system. '
    'These requirements ensure that the project demonstrates essential monitoring and risk '
    'assessment capabilities while remaining achievable within time and resource constraints.'
)

doc.add_heading('3.1 User Needs Analysis', 2)
doc.add_paragraph(
    'The target users of CyberRisk Monitor are students, educators, and small teams who need '
    'an accessible introduction to cybersecurity monitoring concepts. User needs include:'
)

user_needs = [
    ('Transparency', 'Users need to understand why alerts are generated and how risk scores are calculated'),
    ('Simplicity', 'Users require an intuitive interface that does not require extensive cybersecurity expertise'),
    ('Education', 'Users benefit from clear explanations of detected security patterns'),
    ('Flexibility', 'Users need the ability to adjust detection thresholds based on their environment'),
    ('Accessibility', 'Users require cross-platform support and minimal installation complexity'),
]

needs_table = doc.add_table(rows=6, cols=2)
set_table_borders(needs_table)
needs_table.style = 'Table Grid'
needs_table.cell(0, 0).text = 'User Need'
needs_table.cell(0, 1).text = 'Description'
for cell in needs_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

for i, (need, desc) in enumerate(user_needs, 1):
    needs_table.cell(i, 0).text = need
    needs_table.cell(i, 1).text = desc

doc.add_paragraph()
cap = doc.add_paragraph('Table 3.1: User Needs Analysis')
cap.runs[0].italic = True

doc.add_heading('3.2 Functional Requirements', 2)
doc.add_paragraph(
    'Functional requirements define the specific behaviors and capabilities the system must exhibit:'
)

fr_table = doc.add_table(rows=9, cols=4)
set_table_borders(fr_table)
fr_table.style = 'Table Grid'
fr_table.cell(0, 0).text = 'Req ID'
fr_table.cell(0, 1).text = 'Requirement Description'
fr_table.cell(0, 2).text = 'Priority'
fr_table.cell(0, 3).text = 'Verification Method'
for cell in fr_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

fr_data = [
    ('FR-01', 'The system shall parse security logs in syslog-style text, JSON, and CSV formats', 'High', 'Unit Testing'),
    ('FR-02', 'The system shall apply predefined and configurable detection rules to parsed log data', 'High', 'Integration Testing'),
    ('FR-03', 'The system shall identify common security patterns (failed logins, port scanning, privilege escalation)', 'High', 'Functional Testing'),
    ('FR-04', 'The system shall calculate numerical risk scores for detected events', 'High', 'Unit Testing'),
    ('FR-05', 'The system shall categorize detected risks as Low, Medium, High, or Critical', 'High', 'Validation Testing'),
    ('FR-06', 'The system shall generate summary reports of detected findings', 'Medium', 'Functional Testing'),
    ('FR-07', 'The system shall present results through a dashboard or structured report interface', 'Medium', 'User Acceptance'),
    ('FR-08', 'The system shall allow users to adjust detection rule thresholds', 'Low', 'Configuration Testing'),
]

for i, (req_id, desc, priority, verify) in enumerate(fr_data, 1):
    fr_table.cell(i, 0).text = req_id
    fr_table.cell(i, 1).text = desc
    fr_table.cell(i, 2).text = priority
    fr_table.cell(i, 3).text = verify

doc.add_paragraph()
cap = doc.add_paragraph('Table 3.2: Functional Requirements')
cap.runs[0].italic = True

doc.add_heading('3.3 Non-Functional Requirements', 2)
doc.add_paragraph(
    'Non-functional requirements define the quality attributes and constraints on system behavior:'
)

nfr_table = doc.add_table(rows=7, cols=4)
set_table_borders(nfr_table)
nfr_table.style = 'Table Grid'
nfr_table.cell(0, 0).text = 'Req ID'
nfr_table.cell(0, 1).text = 'Requirement Description'
nfr_table.cell(0, 2).text = 'Category'
nfr_table.cell(0, 3).text = 'Acceptance Criteria'
for cell in nfr_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

nfr_data = [
    ('NFR-01', 'Process up to 10,000 log entries within 30 seconds on standard hardware', 'Performance', 'Benchmark testing shows <30s processing time'),
    ('NFR-02', 'User interface understandable for users with basic technical knowledge', 'Usability', 'User can perform analysis without documentation'),
    ('NFR-03', 'Include error handling and internal logging for debugging', 'Reliability', 'No unhandled exceptions; errors logged'),
    ('NFR-04', 'Codebase documented and follows consistent coding standards', 'Maintainability', 'Code review approval; docstrings present'),
    ('NFR-05', 'Cross-platform compatibility (Windows, macOS, Linux)', 'Portability', 'Successful execution on all three platforms'),
    ('NFR-06', 'Modular architecture supporting future enhancements', 'Extensibility', 'Clear separation of concerns in design'),
]

for i, (req_id, desc, cat, criteria) in enumerate(nfr_data, 1):
    nfr_table.cell(i, 0).text = req_id
    nfr_table.cell(i, 1).text = desc
    nfr_table.cell(i, 2).text = cat
    nfr_table.cell(i, 3).text = criteria

doc.add_paragraph()
cap = doc.add_paragraph('Table 3.3: Non-Functional Requirements')
cap.runs[0].italic = True

doc.add_heading('3.4 Technical Requirements', 2)

tech_reqs = [
    ('Programming Language', 'Python 3.10 or later'),
    ('Operating System', 'Cross-platform (Windows, macOS, Linux)'),
    ('Version Control', 'Git with shared GitHub repository'),
    ('Core Libraries', 'Standard Python libraries (json, csv, re, datetime, logging)'),
    ('Optional Frameworks', 'Flask or Streamlit for web-based visualization'),
    ('Configuration Format', 'JSON or YAML files for rules and settings'),
    ('Data Storage', 'File-based; optional SQLite for persistence'),
    ('Testing Framework', 'pytest for unit and integration testing'),
]

tech_table = doc.add_table(rows=9, cols=2)
set_table_borders(tech_table)
tech_table.style = 'Table Grid'
tech_table.cell(0, 0).text = 'Component'
tech_table.cell(0, 1).text = 'Specification'
for cell in tech_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

for i, (comp, spec) in enumerate(tech_reqs, 1):
    tech_table.cell(i, 0).text = comp
    tech_table.cell(i, 1).text = spec

doc.add_paragraph()
cap = doc.add_paragraph('Table 3.4: Technical Requirements')
cap.runs[0].italic = True

# ============================================
# 4. PROJECT METHODOLOGY
# ============================================
doc.add_page_break()
doc.add_heading('4. Project Methodology', 1)

doc.add_heading('4.1 Methodology Selection: Agile-Inspired Iterative Development', 2)
doc.add_paragraph(
    'The CyberRisk Monitor project follows an Agile-inspired iterative development methodology '
    'adapted to the structure of the CMSC 495 capstone course. This approach organizes development '
    'activities into short, iterative cycles aligned with weekly course milestones, allowing for '
    'incremental progress and regular reassessment of project status.'
)

doc.add_heading('4.2 Justification for Methodology Selection', 2)
doc.add_paragraph('The Agile-inspired approach was selected based on the following project characteristics:')

justification_table = doc.add_table(rows=7, cols=2)
set_table_borders(justification_table)
justification_table.style = 'Table Grid'
justification_table.cell(0, 0).text = 'Project Characteristic'
justification_table.cell(0, 1).text = 'How Agile Addresses It'
for cell in justification_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

justification_data = [
    ('Fixed Timeline', 'Weekly iterations ensure continuous progress within the 6-week constraint'),
    ('Small Team', '4-person team benefits from lightweight processes over heavy documentation'),
    ('Educational Context', 'Iterative demos allow instructor feedback and course alignment'),
    ('Technical Uncertainty', 'Flexibility to adapt as team learns cybersecurity concepts'),
    ('Integration Complexity', 'Early and continuous integration reduces end-of-project risk'),
    ('Deliverable Focus', 'Each iteration produces working software demonstrating progress'),
]

for i, (char, addr) in enumerate(justification_data, 1):
    justification_table.cell(i, 0).text = char
    justification_table.cell(i, 1).text = addr

doc.add_paragraph()
cap = doc.add_paragraph('Table 4.1: Methodology Justification')
cap.runs[0].italic = True

doc.add_heading('4.3 Iteration Structure', 2)
doc.add_paragraph(
    'Each weekly iteration follows a structured pattern that balances development work with '
    'team coordination and quality assurance:'
)

iteration_items = [
    'Planning (Monday): Review tasks for the week, assign responsibilities, identify blockers',
    'Development (Monday-Friday): Execute assigned tasks, implement features, write tests',
    'Integration (Thursday): Merge code, resolve conflicts, verify component compatibility',
    'Testing (Thursday-Friday): Run unit tests, integration tests, and manual verification',
    'Review (Friday): Team meeting to demonstrate progress, discuss issues, plan next iteration',
    'Documentation (Ongoing): Update code comments, maintain user documentation'
]
for item in iteration_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 Communication and Collaboration', 2)
doc.add_paragraph('The team employs the following communication practices:')

comm_items = [
    'Weekly team meetings (virtual) for planning and review',
    'Shared GitHub repository for code collaboration and version control',
    'Issue tracking via GitHub Issues for task management',
    'Direct messaging (Discord/Slack) for quick questions and coordination',
    'Shared documentation repository for design documents and notes'
]
for item in comm_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.5 Comparison with Waterfall Approach', 2)
doc.add_paragraph(
    'A traditional Waterfall approach was considered but rejected due to its limitations '
    'in this context. The following table compares both approaches:'
)

comparison_table = doc.add_table(rows=5, cols=3)
set_table_borders(comparison_table)
comparison_table.style = 'Table Grid'
comparison_table.cell(0, 0).text = 'Aspect'
comparison_table.cell(0, 1).text = 'Waterfall'
comparison_table.cell(0, 2).text = 'Agile (Selected)'
for cell in comparison_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

comparison_data = [
    ('Feedback Timing', 'Late (after full implementation)', 'Continuous (each iteration)'),
    ('Risk of Failure', 'High (issues found late)', 'Low (issues found early)'),
    ('Flexibility', 'Low (changes are costly)', 'High (changes are expected)'),
    ('Progress Visibility', 'Limited until end', 'Demonstrated weekly'),
]

for i, (aspect, waterfall, agile) in enumerate(comparison_data, 1):
    comparison_table.cell(i, 0).text = aspect
    comparison_table.cell(i, 1).text = waterfall
    comparison_table.cell(i, 2).text = agile

doc.add_paragraph()
cap = doc.add_paragraph('Table 4.2: Methodology Comparison')
cap.runs[0].italic = True

# ============================================
# 5. WORK BREAKDOWN STRUCTURE AND PROJECT TASKS
# ============================================
doc.add_page_break()
doc.add_heading('5. Work Breakdown Structure and Project Tasks', 1)

doc.add_paragraph(
    'The project work is decomposed into manageable tasks organized by development phase '
    'and functional responsibility. This work breakdown structure supports parallel development '
    'while maintaining clear dependencies between system components.'
)

doc.add_heading('5.1 Work Breakdown Structure Diagram', 2)
doc.add_paragraph(
    'The following diagram illustrates the hierarchical decomposition of project work:'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\wbs_diagram.png',
                       'Figure 5.1: Work Breakdown Structure (WBS) Diagram', 6.5)

doc.add_heading('5.2 Task Breakdown by Phase', 2)

doc.add_heading('Phase 1: Planning and Design (Week 3)', 3)
phase1_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase1_table)
phase1_table.style = 'Table Grid'
phase1_table.cell(0, 0).text = 'Task ID'
phase1_table.cell(0, 1).text = 'Task Description'
phase1_table.cell(0, 2).text = 'Owner'
phase1_table.cell(0, 3).text = 'Est. Hours'
phase1_table.cell(0, 4).text = 'Deliverable'
for cell in phase1_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase1_data = [
    ('T1.1', 'Finalize requirements documentation', 'Mustafa', '6', 'Requirements Document'),
    ('T1.2', 'Design system architecture', 'Nicholas', '8', 'Architecture Diagram'),
    ('T1.3', 'Define detection rules catalog', 'Daniel', '6', 'Rules Specification'),
    ('T1.4', 'Set up development environment', 'All', '4', 'Working Dev Environment'),
]

for i, row_data in enumerate(phase1_data, 1):
    for j, cell_data in enumerate(row_data):
        phase1_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('Phase 2: Core Development I (Week 4)', 3)
phase2_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase2_table)
phase2_table.style = 'Table Grid'
phase2_table.cell(0, 0).text = 'Task ID'
phase2_table.cell(0, 1).text = 'Task Description'
phase2_table.cell(0, 2).text = 'Owner'
phase2_table.cell(0, 3).text = 'Est. Hours'
phase2_table.cell(0, 4).text = 'Deliverable'
for cell in phase2_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase2_data = [
    ('T2.1', 'Implement LogParser base module', 'Daniel', '10', 'LogParser.py'),
    ('T2.2', 'Implement syslog format handler', 'Daniel', '6', 'Syslog parser'),
    ('T2.3', 'Implement JSON/CSV format handlers', 'Daniel', '6', 'JSON/CSV parsers'),
    ('T2.4', 'Create unit tests for parsing', 'Daniel', '4', 'test_logparser.py'),
]

for i, row_data in enumerate(phase2_data, 1):
    for j, cell_data in enumerate(row_data):
        phase2_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('Phase 3: Core Development II (Week 5)', 3)
phase3_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase3_table)
phase3_table.style = 'Table Grid'
phase3_table.cell(0, 0).text = 'Task ID'
phase3_table.cell(0, 1).text = 'Task Description'
phase3_table.cell(0, 2).text = 'Owner'
phase3_table.cell(0, 3).text = 'Est. Hours'
phase3_table.cell(0, 4).text = 'Deliverable'
for cell in phase3_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase3_data = [
    ('T3.1', 'Implement RuleEngine module', 'Mustafa', '12', 'RuleEngine.py'),
    ('T3.2', 'Implement RiskScorer module', 'Nicholas', '10', 'RiskScorer.py'),
    ('T3.3', 'Create detection rules (10-15 rules)', 'Mustafa', '8', 'rules_config.json'),
    ('T3.4', 'Unit tests for rules and scoring', 'Nicholas', '6', 'test_rules.py'),
]

for i, row_data in enumerate(phase3_data, 1):
    for j, cell_data in enumerate(row_data):
        phase3_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('Phase 4: Integration (Week 6)', 3)
phase4_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase4_table)
phase4_table.style = 'Table Grid'
phase4_table.cell(0, 0).text = 'Task ID'
phase4_table.cell(0, 1).text = 'Task Description'
phase4_table.cell(0, 2).text = 'Owner'
phase4_table.cell(0, 3).text = 'Est. Hours'
phase4_table.cell(0, 4).text = 'Deliverable'
for cell in phase4_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase4_data = [
    ('T4.1', 'Integrate LogParser with RuleEngine', 'Nicholas', '6', 'Integrated pipeline'),
    ('T4.2', 'Integrate RiskScorer with pipeline', 'Nicholas', '6', 'Full analysis chain'),
    ('T4.3', 'Implement ReportGenerator module', 'Cassandra', '10', 'ReportGenerator.py'),
    ('T4.4', 'Integration testing', 'All', '8', 'Integration test suite'),
]

for i, row_data in enumerate(phase4_data, 1):
    for j, cell_data in enumerate(row_data):
        phase4_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('Phase 5: UI and Testing (Week 7)', 3)
phase5_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase5_table)
phase5_table.style = 'Table Grid'
phase5_table.cell(0, 0).text = 'Task ID'
phase5_table.cell(0, 1).text = 'Task Description'
phase5_table.cell(0, 2).text = 'Owner'
phase5_table.cell(0, 3).text = 'Est. Hours'
phase5_table.cell(0, 4).text = 'Deliverable'
for cell in phase5_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase5_data = [
    ('T5.1', 'Implement UserInterface (CLI)', 'Cassandra', '8', 'cli.py'),
    ('T5.2', 'Implement dashboard visualization', 'Cassandra', '10', 'Dashboard module'),
    ('T5.3', 'System testing and bug fixes', 'All', '10', 'Test reports'),
    ('T5.4', 'Performance optimization', 'Mustafa', '4', 'Optimized code'),
]

for i, row_data in enumerate(phase5_data, 1):
    for j, cell_data in enumerate(row_data):
        phase5_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('Phase 6: Final Delivery (Week 8)', 3)
phase6_table = doc.add_table(rows=5, cols=5)
set_table_borders(phase6_table)
phase6_table.style = 'Table Grid'
phase6_table.cell(0, 0).text = 'Task ID'
phase6_table.cell(0, 1).text = 'Task Description'
phase6_table.cell(0, 2).text = 'Owner'
phase6_table.cell(0, 3).text = 'Est. Hours'
phase6_table.cell(0, 4).text = 'Deliverable'
for cell in phase6_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

phase6_data = [
    ('T6.1', 'Finalize code documentation', 'All', '6', 'Documented codebase'),
    ('T6.2', 'Create user manual', 'Cassandra', '6', 'User Manual'),
    ('T6.3', 'Performance validation', 'Daniel', '4', 'Benchmark results'),
    ('T6.4', 'Prepare final demonstration', 'Mustafa', '6', 'Demo presentation'),
]

for i, row_data in enumerate(phase6_data, 1):
    for j, cell_data in enumerate(row_data):
        phase6_table.cell(i, j).text = cell_data

doc.add_paragraph()

doc.add_heading('5.3 Task Dependencies', 2)
doc.add_paragraph('The following critical dependencies exist between tasks:')

dependencies = [
    'T2.1-T2.4 (LogParser) must complete before T4.1 (Integration with RuleEngine)',
    'T3.1 (RuleEngine) must complete before T4.1 (Integration)',
    'T3.2 (RiskScorer) must complete before T4.2 (Pipeline Integration)',
    'T4.1-T4.2 (Integration) must complete before T5.1-T5.2 (UI Development)',
    'T4.3 (ReportGenerator) must complete before T5.2 (Dashboard)',
    'All development tasks must complete before T6.3 (Performance Validation)',
    'T6.1-T6.3 must complete before T6.4 (Final Demonstration)'
]
for dep in dependencies:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_heading('5.4 Total Effort Summary', 2)
summary_table = doc.add_table(rows=8, cols=3)
set_table_borders(summary_table)
summary_table.style = 'Table Grid'
summary_table.cell(0, 0).text = 'Phase'
summary_table.cell(0, 1).text = 'Total Hours'
summary_table.cell(0, 2).text = 'Team Members Involved'
for cell in summary_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

summary_data = [
    ('Phase 1: Planning & Design', '24', 'All (6 hrs each)'),
    ('Phase 2: Core Development I', '26', 'Daniel (primary)'),
    ('Phase 3: Core Development II', '36', 'Mustafa, Nicholas'),
    ('Phase 4: Integration', '30', 'Nicholas, Cassandra'),
    ('Phase 5: UI & Testing', '32', 'Cassandra, All'),
    ('Phase 6: Final Delivery', '22', 'All'),
    ('TOTAL', '170', '~42.5 hrs per member'),
]

for i, row_data in enumerate(summary_data, 1):
    for j, cell_data in enumerate(row_data):
        summary_table.cell(i, j).text = cell_data
        if i == 7:  # Bold the total row
            summary_table.cell(i, j).paragraphs[0].runs[0].bold = True

doc.add_paragraph()
cap = doc.add_paragraph('Table 5.1: Total Effort Summary by Phase')
cap.runs[0].italic = True

# ============================================
# 6. PROJECT SCHEDULE
# ============================================
doc.add_page_break()
doc.add_heading('6. Project Schedule', 1)

doc.add_paragraph(
    'The project schedule spans six weeks, corresponding to Weeks 3 through 8 of the capstone '
    'course. Milestones are structured to ensure consistent progress while allocating sufficient '
    'time for integration, testing, and refinement.'
)

doc.add_heading('6.1 Project Timeline (Gantt Chart)', 2)
doc.add_paragraph(
    'The following Gantt chart illustrates the project timeline with task durations, dependencies, '
    'and milestones:'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\gantt_chart.png',
                       'Figure 6.1: Project Schedule Gantt Chart', 6.5)

doc.add_heading('6.2 Milestones and Deadlines', 2)

milestones_table = doc.add_table(rows=7, cols=4)
set_table_borders(milestones_table)
milestones_table.style = 'Table Grid'
milestones_table.cell(0, 0).text = 'Milestone'
milestones_table.cell(0, 1).text = 'Target Date'
milestones_table.cell(0, 2).text = 'Deliverables'
milestones_table.cell(0, 3).text = 'Success Criteria'
for cell in milestones_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

milestone_data = [
    ('M1: Design Complete', 'End of Week 3', 'Design document, Architecture diagram', 'Approved by team lead'),
    ('M2: Parser Ready', 'End of Week 4', 'LogParser module, Unit tests', 'All format handlers tested'),
    ('M3: Core Engine Ready', 'End of Week 5', 'RuleEngine, RiskScorer modules', 'Rules execute correctly'),
    ('M4: Integration Complete', 'End of Week 6', 'Integrated pipeline, ReportGenerator', 'End-to-end test passes'),
    ('M5: UI Complete', 'End of Week 7', 'UserInterface, Dashboard', 'User can run full analysis'),
    ('M6: Project Delivery', 'End of Week 8', 'Final code, Documentation, Demo', 'All requirements met'),
]

for i, row_data in enumerate(milestone_data, 1):
    for j, cell_data in enumerate(row_data):
        milestones_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 6.1: Project Milestones and Deadlines')
cap.runs[0].italic = True

doc.add_heading('6.3 Weekly Schedule Detail', 2)

weekly_table = doc.add_table(rows=7, cols=3)
set_table_borders(weekly_table)
weekly_table.style = 'Table Grid'
weekly_table.cell(0, 0).text = 'Week'
weekly_table.cell(0, 1).text = 'Primary Focus'
weekly_table.cell(0, 2).text = 'Key Activities'
for cell in weekly_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

weekly_data = [
    ('Week 3', 'Planning & Design', 'Requirements finalization, Architecture design, Rule definition, Environment setup'),
    ('Week 4', 'Core Development I', 'LogParser implementation, Format handlers (syslog, JSON, CSV), Unit testing'),
    ('Week 5', 'Core Development II', 'RuleEngine implementation, RiskScorer implementation, Detection rules creation'),
    ('Week 6', 'Integration', 'Component integration, ReportGenerator development, Integration testing'),
    ('Week 7', 'UI & Testing', 'CLI and Dashboard development, System testing, Performance optimization'),
    ('Week 8', 'Final Delivery', 'Documentation, Performance validation, Demo preparation, Final submission'),
]

for i, row_data in enumerate(weekly_data, 1):
    for j, cell_data in enumerate(row_data):
        weekly_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 6.2: Weekly Schedule Detail')
cap.runs[0].italic = True

doc.add_heading('6.4 Critical Path', 2)
doc.add_paragraph(
    'The critical path for this project runs through the following sequence of dependent tasks:'
)

critical_path = [
    'Requirements Analysis (Week 3)',
    'LogParser Module (Week 4)',
    'RuleEngine Module (Week 5)',
    'Component Integration (Week 6)',
    'UserInterface Development (Week 7)',
    'Final Demo Preparation (Week 8)'
]

doc.add_paragraph('Critical Path: ' + ' -> '.join(critical_path))

doc.add_paragraph(
    'Any delays in these tasks will directly impact the final delivery date. Buffer time has been '
    'allocated within each phase to absorb minor delays without affecting subsequent milestones.'
)

# ============================================
# 7. PROJECT RESOURCES
# ============================================
doc.add_page_break()
doc.add_heading('7. Project Resources', 1)

doc.add_paragraph(
    'This section identifies and allocates the personnel, equipment, and budget resources '
    'required for successful project completion.'
)

doc.add_heading('7.1 Team Member Allocation', 2)

team_table = doc.add_table(rows=5, cols=5)
set_table_borders(team_table)
team_table.style = 'Table Grid'
team_table.cell(0, 0).text = 'Team Member'
team_table.cell(0, 1).text = 'Role'
team_table.cell(0, 2).text = 'Primary Responsibilities'
team_table.cell(0, 3).text = 'Availability'
team_table.cell(0, 4).text = 'Est. Hours'
for cell in team_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

team_data = [
    ('Mustafa Black-Castle', 'Project Lead', 'RuleEngine, Coordination, Demo', '10-12 hrs/week', '40'),
    ('Daniel S. Garrett', 'Lead Developer', 'LogParser, Testing Lead', '10-12 hrs/week', '40'),
    ('Nicholas Porpora', 'Developer', 'RiskScorer, Integration', '10-12 hrs/week', '40'),
    ('Cassandra Santacruz', 'Developer', 'UI, Documentation, Reports', '10-12 hrs/week', '40'),
]

for i, row_data in enumerate(team_data, 1):
    for j, cell_data in enumerate(row_data):
        team_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 7.1: Team Member Allocation')
cap.runs[0].italic = True

doc.add_heading('7.2 Resource Allocation Visualization', 2)
add_image_with_caption(doc, r'E:\CyberRisk\docs\images\resource_allocation.png',
                       'Figure 7.1: Team Workload and Weekly Effort Distribution', 6.5)

doc.add_heading('7.3 Equipment and Software Resources', 2)

equip_table = doc.add_table(rows=9, cols=4)
set_table_borders(equip_table)
equip_table.style = 'Table Grid'
equip_table.cell(0, 0).text = 'Resource'
equip_table.cell(0, 1).text = 'Type'
equip_table.cell(0, 2).text = 'Cost'
equip_table.cell(0, 3).text = 'Availability'
for cell in equip_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

equip_data = [
    ('Personal Computers', 'Hardware', '$0 (owned)', 'All team members'),
    ('Python 3.10+', 'Software', '$0 (free)', 'Available'),
    ('Git/GitHub', 'Version Control', '$0 (free tier)', 'Available'),
    ('VS Code/PyCharm', 'IDE', '$0 (free editions)', 'Available'),
    ('pytest', 'Testing Framework', '$0 (open source)', 'Available'),
    ('Streamlit/Flask', 'Web Framework', '$0 (open source)', 'Available'),
    ('Discord/Slack', 'Communication', '$0 (free tier)', 'Available'),
    ('Internet Access', 'Infrastructure', 'Personal expense', 'All team members'),
]

for i, row_data in enumerate(equip_data, 1):
    for j, cell_data in enumerate(row_data):
        equip_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 7.2: Equipment and Software Resources')
cap.runs[0].italic = True

doc.add_heading('7.4 Budget Summary', 2)
doc.add_paragraph(
    'This project operates with a zero-dollar budget. All tools, frameworks, and resources '
    'utilized are free, open-source, or already owned by team members. The only implicit costs '
    'are team member time (estimated at 160 total person-hours) and personal internet access.'
)

budget_table = doc.add_table(rows=5, cols=2)
set_table_borders(budget_table)
budget_table.style = 'Table Grid'
budget_table.cell(0, 0).text = 'Cost Category'
budget_table.cell(0, 1).text = 'Amount'
for cell in budget_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

budget_data = [
    ('Software Licenses', '$0'),
    ('Hardware', '$0 (existing equipment)'),
    ('Cloud Services', '$0 (not required)'),
    ('TOTAL PROJECT BUDGET', '$0'),
]

for i, row_data in enumerate(budget_data, 1):
    budget_table.cell(i, 0).text = row_data[0]
    budget_table.cell(i, 1).text = row_data[1]

doc.add_paragraph()
cap = doc.add_paragraph('Table 7.3: Budget Summary')
cap.runs[0].italic = True

# ============================================
# 8. PROJECT RISKS
# ============================================
doc.add_page_break()
doc.add_heading('8. Project Risks', 1)

doc.add_paragraph(
    'Cybersecurity projects often involve technical complexity and integration challenges, '
    'making proactive risk management essential. This section identifies potential risks, '
    'assesses their impact and likelihood, and outlines mitigation strategies.'
)

doc.add_heading('8.1 Risk Assessment Matrix', 2)
doc.add_paragraph(
    'The following matrix visualizes identified risks based on their likelihood and potential impact:'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\risk_matrix.png',
                       'Figure 8.1: Risk Assessment Matrix', 6)

doc.add_heading('8.2 Risk Register', 2)

risk_table = doc.add_table(rows=7, cols=6)
set_table_borders(risk_table)
risk_table.style = 'Table Grid'
risk_table.cell(0, 0).text = 'Risk ID'
risk_table.cell(0, 1).text = 'Risk Description'
risk_table.cell(0, 2).text = 'Likelihood'
risk_table.cell(0, 3).text = 'Impact'
risk_table.cell(0, 4).text = 'Risk Level'
risk_table.cell(0, 5).text = 'Mitigation Strategy'
for cell in risk_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

risk_data = [
    ('R1', 'Team member unavailability due to other commitments', 'Possible', 'Major', 'High', 'Cross-train on modules; maintain documentation; buffer time in schedule'),
    ('R2', 'Technical complexity exceeds team expertise', 'Unlikely', 'Moderate', 'Medium', 'Early research; leverage online resources; simplify scope if needed'),
    ('R3', 'Scope creep from feature requests', 'Possible', 'Moderate', 'Medium', 'Strict scope boundary enforcement; change control process'),
    ('R4', 'Integration issues between modules', 'Unlikely', 'Major', 'Medium', 'Define clear interfaces early; continuous integration testing'),
    ('R5', 'Learning curve for cybersecurity concepts', 'Possible', 'Minor', 'Medium', 'Early research; focus on educational resources; team knowledge sharing'),
    ('R6', 'Poor quality or insufficient test data', 'Rare', 'Moderate', 'Low', 'Create synthetic datasets early; validate data quality before testing'),
]

for i, row_data in enumerate(risk_data, 1):
    for j, cell_data in enumerate(row_data):
        risk_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 8.1: Risk Register')
cap.runs[0].italic = True

doc.add_heading('8.3 Detailed Mitigation Plans', 2)

doc.add_heading('R1: Team Member Unavailability', 3)
doc.add_paragraph(
    'Impact: Delays in task completion; potential knowledge gaps. '
    'Mitigation Actions:'
)
r1_mitigations = [
    'Maintain comprehensive documentation for all modules',
    'Ensure at least two team members understand each critical component',
    'Hold weekly sync meetings to share progress and knowledge',
    'Build buffer time into schedule for unexpected absences',
    'Use GitHub Issues to track all tasks with clear descriptions'
]
for m in r1_mitigations:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('R2: Technical Complexity', 3)
doc.add_paragraph(
    'Impact: Extended development time; potential feature reduction. '
    'Mitigation Actions:'
)
r2_mitigations = [
    'Conduct early research during Week 3 planning phase',
    'Identify online tutorials and resources for unfamiliar concepts',
    'Start with simple implementations and iterate',
    'Seek instructor guidance when facing significant blockers',
    'Be prepared to simplify scope while maintaining core functionality'
]
for m in r2_mitigations:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('R3: Scope Creep', 3)
doc.add_paragraph(
    'Impact: Delayed delivery; incomplete core features. '
    'Mitigation Actions:'
)
r3_mitigations = [
    'Document clear scope boundaries in this design document',
    'Require team consensus for any scope changes',
    'Maintain a "future enhancements" list for ideas outside current scope',
    'Prioritize core functionality over nice-to-have features',
    'Review scope weekly during team meetings'
]
for m in r3_mitigations:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('R4: Integration Issues', 3)
doc.add_paragraph(
    'Impact: System fails to work as a cohesive unit. '
    'Mitigation Actions:'
)
r4_mitigations = [
    'Define clear interfaces between modules during design phase',
    'Use consistent data structures across all components',
    'Implement continuous integration from Week 4 onwards',
    'Write integration tests before full integration begins',
    'Conduct code reviews to ensure interface compliance'
]
for m in r4_mitigations:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('8.4 Contingency Plans', 2)

contingency_table = doc.add_table(rows=4, cols=3)
set_table_borders(contingency_table)
contingency_table.style = 'Table Grid'
contingency_table.cell(0, 0).text = 'Scenario'
contingency_table.cell(0, 1).text = 'Trigger Condition'
contingency_table.cell(0, 2).text = 'Contingency Action'
for cell in contingency_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

contingency_data = [
    ('Major delay (>1 week)', 'Milestone missed by >3 days', 'Reduce detection rules to 8; simplify UI to CLI only'),
    ('Team member dropout', 'Member unavailable for >1 week', 'Redistribute tasks; focus on critical path items'),
    ('Technical blocker', 'Issue unresolved for >3 days', 'Seek external help; simplify affected component'),
]

for i, row_data in enumerate(contingency_data, 1):
    for j, cell_data in enumerate(row_data):
        contingency_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 8.2: Contingency Plans')
cap.runs[0].italic = True

# ============================================
# 9. PROJECT EVALUATION PLAN
# ============================================
doc.add_page_break()
doc.add_heading('9. Project Evaluation Plan', 1)

doc.add_paragraph(
    'Project success will be evaluated using both qualitative and quantitative criteria. '
    'This section establishes clear, measurable evaluation criteria aligned with project goals.'
)

doc.add_heading('9.1 Evaluation Criteria', 2)

eval_table = doc.add_table(rows=9, cols=4)
set_table_borders(eval_table)
eval_table.style = 'Table Grid'
eval_table.cell(0, 0).text = 'Criterion'
eval_table.cell(0, 1).text = 'Metric'
eval_table.cell(0, 2).text = 'Target'
eval_table.cell(0, 3).text = 'Measurement Method'
for cell in eval_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

eval_data = [
    ('Functional Completeness', 'Requirements implemented', '100% of FR-01 to FR-08', 'Requirements traceability matrix'),
    ('Detection Accuracy', 'Rules correctly identify patterns', '>90% true positive rate', 'Testing with known datasets'),
    ('Performance', 'Log processing speed', '<30 sec for 10,000 entries', 'Benchmark testing'),
    ('Code Quality', 'Test coverage', '>80% code coverage', 'pytest-cov report'),
    ('Documentation', 'Completeness', 'All modules documented', 'Documentation review'),
    ('Usability', 'User can complete analysis', 'Without documentation reference', 'User acceptance testing'),
    ('Maintainability', 'Code follows standards', 'Passes linting checks', 'Automated code analysis'),
    ('Milestone Adherence', 'On-time delivery', '100% milestones met', 'Schedule tracking'),
]

for i, row_data in enumerate(eval_data, 1):
    for j, cell_data in enumerate(row_data):
        eval_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 9.1: Evaluation Criteria and Metrics')
cap.runs[0].italic = True

doc.add_heading('9.2 Progress Monitoring', 2)
doc.add_paragraph('Progress will be monitored throughout the project using the following methods:')

monitoring = [
    ('Weekly Status Reviews', 'Team meeting every Friday to assess progress against milestones'),
    ('GitHub Metrics', 'Track commits, pull requests, and issue closure rates'),
    ('Test Results', 'Monitor unit and integration test pass rates'),
    ('Code Reviews', 'Peer review of all code before merging to main branch'),
    ('Milestone Checkpoints', 'Formal review at each milestone completion'),
]

for method, desc in monitoring:
    p = doc.add_paragraph()
    p.add_run(method + ': ').bold = True
    p.add_run(desc)

doc.add_heading('9.3 Quality Assurance Checkpoints', 2)

qa_table = doc.add_table(rows=7, cols=3)
set_table_borders(qa_table)
qa_table.style = 'Table Grid'
qa_table.cell(0, 0).text = 'Checkpoint'
qa_table.cell(0, 1).text = 'Week'
qa_table.cell(0, 2).text = 'Quality Gate Criteria'
for cell in qa_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

qa_data = [
    ('Design Review', '3', 'Architecture approved; requirements documented'),
    ('Parser Quality Gate', '4', 'All format tests pass; code reviewed'),
    ('Engine Quality Gate', '5', 'Detection rules validated; scoring correct'),
    ('Integration Quality Gate', '6', 'End-to-end test passes; no critical bugs'),
    ('UI Quality Gate', '7', 'User acceptance criteria met'),
    ('Final Quality Gate', '8', 'All criteria met; demo ready'),
]

for i, row_data in enumerate(qa_data, 1):
    for j, cell_data in enumerate(row_data):
        qa_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 9.2: Quality Assurance Checkpoints')
cap.runs[0].italic = True

doc.add_heading('9.4 Success Definition', 2)
doc.add_paragraph('The project will be considered successful if:')

success_criteria = [
    'All functional requirements (FR-01 through FR-08) are implemented and working',
    'The system can process 10,000 log entries within 30 seconds',
    'At least 10 detection rules are implemented and validated',
    'The user interface allows complete analysis workflow without errors',
    'Code documentation covers all public interfaces',
    'Final demonstration successfully shows all capabilities',
    'All six milestones are completed by their target dates'
]
for criterion in success_criteria:
    doc.add_paragraph(criterion, style='List Bullet')

doc.add_heading('9.5 Post-Project Evaluation', 2)
doc.add_paragraph(
    'Following project completion, a retrospective evaluation will assess:'
)

retrospective = [
    'What worked well in the development process',
    'What challenges were encountered and how they were addressed',
    'Lessons learned for future projects',
    'Potential improvements if the project were to continue',
    'Team collaboration effectiveness'
]
for item in retrospective:
    doc.add_paragraph(item, style='List Bullet')

# ============================================
# 10. SYSTEM DESIGN
# ============================================
doc.add_page_break()
doc.add_heading('10. System Design', 1)

doc.add_heading('10.1 System Architecture Overview', 2)
doc.add_paragraph(
    'CyberRisk Monitor follows a modular system architecture that separates data ingestion, '
    'rule evaluation, risk assessment, and presentation. This separation of concerns supports '
    'maintainability, testability, and clear responsibility boundaries.'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\system_architecture.png',
                       'Figure 10.1: System Architecture and Data Flow Diagram', 6.5)

doc.add_heading('10.2 Component Descriptions', 2)

components = [
    ('LogParser', 'Ingests and normalizes log data from multiple input formats (syslog, JSON, CSV). Responsible for format detection, parsing, validation, and conversion to a unified internal representation.'),
    ('RuleEngine', 'Applies detection rules to parsed log entries. Loads rule configurations, evaluates conditions against log data, and generates detection events for matching patterns.'),
    ('RiskScorer', 'Calculates risk scores and assigns severity levels. Implements the scoring algorithm based on rule severity, event frequency, and contextual factors.'),
    ('ReportGenerator', 'Produces summary reports and visualizations. Aggregates detection results, calculates statistics, and formats output for display or export.'),
    ('UserInterface', 'Provides user interaction and output display. Supports command-line interface and optional graphical dashboard for analysis workflow.'),
]

for comp_name, comp_desc in components:
    p = doc.add_paragraph()
    p.add_run(comp_name + ': ').bold = True
    p.add_run(comp_desc)

doc.add_heading('10.3 UML Class Diagram', 2)
doc.add_paragraph(
    'The following detailed UML class diagram illustrates the object-oriented design of the system:'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\class_diagram_detailed.png',
                       'Figure 10.2: Detailed UML Class Diagram', 6.5)

doc.add_heading('10.4 Original Conceptual Class Diagram', 2)
doc.add_paragraph(
    'The following diagram shows the initial conceptual design that guided the detailed implementation:'
)

add_image_with_caption(doc, r'E:\CyberRisk\docs\images\image1.png',
                       'Figure 10.3: Conceptual UML Class Diagram', 5.5)

doc.add_heading('10.5 Data Flow Description', 2)
doc.add_paragraph(
    'The data flow within CyberRisk Monitor follows a deterministic and sequential processing model:'
)

data_flow_steps = [
    'User selects one or more log files through the UserInterface',
    'LogParser identifies the format of each input file and parses entries',
    'Parsed entries are normalized into a unified LogEntry representation',
    'RuleEngine evaluates each LogEntry against all active detection rules',
    'Matching rules generate Detection events with associated metadata',
    'RiskScorer calculates scores for each Detection based on severity and context',
    'Detections are categorized into Low, Medium, High, or Critical severity levels',
    'ReportGenerator aggregates results and generates summary statistics',
    'UserInterface displays results through dashboard or exports reports'
]

for i, step in enumerate(data_flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('10.6 Rule Design', 2)
doc.add_paragraph(
    'Detection rules are designed to identify common security patterns. Each rule consists of:'
)

rule_elements = [
    ('Rule ID', 'Unique identifier for tracking and reference'),
    ('Name', 'Human-readable description of the detection pattern'),
    ('Pattern', 'Regular expression or condition to match against log entries'),
    ('Threshold', 'Number of occurrences required to trigger detection'),
    ('Time Window', 'Duration over which threshold is evaluated'),
    ('Severity', 'Impact level (1-10) assigned to matched events'),
    ('Enabled', 'Flag to enable/disable the rule'),
]

for element, desc in rule_elements:
    p = doc.add_paragraph()
    p.add_run(element + ': ').bold = True
    p.add_run(desc)

doc.add_heading('10.7 Risk Scoring Model', 2)
doc.add_paragraph(
    'The risk scoring model aggregates individual rule detections into cumulative risk scores:'
)

scoring_formula = [
    'Base Score = Rule Severity (1-10)',
    'Frequency Multiplier = log2(occurrence count + 1)',
    'Event Score = Base Score * Frequency Multiplier',
    'Total Risk Score = Sum of all Event Scores',
    'Risk Level = Categorized based on thresholds (Low: 0-25, Medium: 26-50, High: 51-75, Critical: 76+)'
]

for formula in scoring_formula:
    doc.add_paragraph(formula, style='List Bullet')

# ============================================
# 11. TESTING AND VALIDATION STRATEGY
# ============================================
doc.add_page_break()
doc.add_heading('11. Testing and Validation Strategy', 1)

doc.add_paragraph(
    'Testing and validation are conducted throughout the development lifecycle to ensure '
    'functional correctness, reliability, and alignment with project requirements.'
)

doc.add_heading('11.1 Testing Levels', 2)

testing_table = doc.add_table(rows=5, cols=3)
set_table_borders(testing_table)
testing_table.style = 'Table Grid'
testing_table.cell(0, 0).text = 'Testing Level'
testing_table.cell(0, 1).text = 'Scope'
testing_table.cell(0, 2).text = 'Tools/Approach'
for cell in testing_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

testing_data = [
    ('Unit Testing', 'Individual functions and methods', 'pytest with assertions'),
    ('Integration Testing', 'Component interactions', 'pytest with test fixtures'),
    ('System Testing', 'Complete end-to-end workflow', 'Manual and automated scenarios'),
    ('Performance Testing', 'Processing speed and efficiency', 'Benchmark with 10K entries'),
]

for i, row_data in enumerate(testing_data, 1):
    for j, cell_data in enumerate(row_data):
        testing_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table 11.1: Testing Levels')
cap.runs[0].italic = True

doc.add_heading('11.2 Test Cases by Component', 2)

doc.add_heading('LogParser Test Cases', 3)
parser_tests = [
    'TC-LP-01: Parse valid syslog format file successfully',
    'TC-LP-02: Parse valid JSON format file successfully',
    'TC-LP-03: Parse valid CSV format file successfully',
    'TC-LP-04: Handle malformed log entries gracefully',
    'TC-LP-05: Handle empty files without error',
    'TC-LP-06: Normalize timestamps across formats consistently',
]
for test in parser_tests:
    doc.add_paragraph(test, style='List Bullet')

doc.add_heading('RuleEngine Test Cases', 3)
engine_tests = [
    'TC-RE-01: Load rules from configuration file',
    'TC-RE-02: Detect repeated failed login attempts',
    'TC-RE-03: Detect port scanning patterns',
    'TC-RE-04: Detect privilege escalation attempts',
    'TC-RE-05: Handle rules with different threshold values',
    'TC-RE-06: Enable/disable individual rules',
]
for test in engine_tests:
    doc.add_paragraph(test, style='List Bullet')

doc.add_heading('RiskScorer Test Cases', 3)
scorer_tests = [
    'TC-RS-01: Calculate correct score for single detection',
    'TC-RS-02: Aggregate scores for multiple detections',
    'TC-RS-03: Categorize Low severity correctly (0-25)',
    'TC-RS-04: Categorize Medium severity correctly (26-50)',
    'TC-RS-05: Categorize High severity correctly (51-75)',
    'TC-RS-06: Categorize Critical severity correctly (76+)',
]
for test in scorer_tests:
    doc.add_paragraph(test, style='List Bullet')

doc.add_heading('11.3 Validation Approach', 2)
doc.add_paragraph(
    'Validation testing ensures the system meets defined requirements and user needs:'
)

validation_items = [
    ('Functional Validation', 'Verify all FR-01 through FR-08 requirements are implemented correctly'),
    ('Performance Validation', 'Confirm processing of 10,000 entries completes within 30 seconds'),
    ('Usability Validation', 'Verify users can complete analysis workflow without documentation'),
    ('Accuracy Validation', 'Test detection rules against known attack patterns in sample data'),
]

for item, desc in validation_items:
    p = doc.add_paragraph()
    p.add_run(item + ': ').bold = True
    p.add_run(desc)

doc.add_heading('11.4 Test Data', 2)
doc.add_paragraph(
    'Testing will use synthetic log data containing known patterns:'
)

test_data_items = [
    'Normal baseline logs representing typical system activity',
    'Failed login sequences with varying patterns and frequencies',
    'Port scanning simulation data',
    'Privilege escalation attempt patterns',
    'Mixed datasets combining normal and malicious patterns',
    'Edge cases including empty fields, unusual timestamps, and malformed entries'
]
for item in test_data_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================
# 12. CONCLUSION
# ============================================
doc.add_page_break()
doc.add_heading('12. Conclusion', 1)

doc.add_paragraph(
    'This Project Design document provides a comprehensive roadmap for the successful execution '
    'of the CyberRisk Monitor capstone project. By defining clear scope boundaries, detailed system '
    'requirements, appropriate development methodology, structured task breakdown, realistic scheduling, '
    'thorough risk management strategies, and measurable evaluation criteria, the project establishes '
    'a solid foundation for development, testing, and final delivery.'
)

doc.add_paragraph(
    'The modular architecture design ensures maintainability and supports the educational objectives '
    'of the capstone course. The rule-based detection approach provides transparency and interpretability, '
    'making the system suitable for learning cybersecurity monitoring concepts. The Agile-inspired methodology '
    'enables iterative progress with continuous feedback and adaptation.'
)

doc.add_paragraph(
    'Successful completion of this design and its associated implementation will demonstrate the team\'s '
    'ability to apply computer science principles in a structured, real-world context. CyberRisk Monitor '
    'serves as both a functional monitoring prototype and an educational tool, reinforcing core concepts '
    'such as log analysis, rule-based detection, and risk assessment.'
)

doc.add_paragraph(
    'This design document establishes the clear foundation required for development, testing, and final '
    'demonstration in fulfillment of the CMSC 495 capstone requirements.'
)

# ============================================
# APPENDIX A: DETAILED CLASS SPECIFICATIONS
# ============================================
doc.add_page_break()
doc.add_heading('Appendix A: Detailed Class Specifications', 1)

doc.add_heading('A.1 LogEntry Class', 2)
spec = '''class LogEntry:
    """Represents a normalized log entry from any supported format."""

    Attributes:
        timestamp: datetime - When the event occurred
        source_ip: str - IP address of the event source
        dest_ip: str - IP address of the destination (if applicable)
        event_type: str - Category of the event
        user: str - Username associated with the event
        message: str - Human-readable event description
        raw_data: str - Original unparsed log line
        metadata: dict - Additional format-specific fields

    Methods:
        toDict() -> dict: Convert to dictionary representation
        fromDict(data: dict) -> LogEntry: Create from dictionary
        __str__() -> str: Human-readable string representation'''
doc.add_paragraph(spec)

doc.add_heading('A.2 Rule Class', 2)
spec = '''class Rule:
    """Represents a detection rule for identifying security patterns."""

    Attributes:
        rule_id: str - Unique identifier
        name: str - Human-readable name
        description: str - Detailed explanation
        pattern: str - Regex pattern or condition
        field: str - LogEntry field to match against
        threshold: int - Occurrences needed to trigger
        time_window: int - Seconds for threshold evaluation
        severity: int - Impact level (1-10)
        enabled: bool - Whether rule is active

    Methods:
        matches(entry: LogEntry) -> bool: Check if entry matches rule
        getSeverity() -> int: Return severity level
        enable() -> None: Activate the rule
        disable() -> None: Deactivate the rule'''
doc.add_paragraph(spec)

doc.add_heading('A.3 Detection Class', 2)
spec = '''class Detection:
    """Represents a detected security event."""

    Attributes:
        detection_id: str - Unique identifier
        rule_id: str - ID of triggering rule
        rule_name: str - Name of triggering rule
        log_entries: List[LogEntry] - Entries that triggered detection
        timestamp: datetime - When detection was generated
        score: int - Calculated risk score
        severity: str - Categorized severity level

    Methods:
        toDict() -> dict: Convert to dictionary
        getSeverity() -> str: Return severity category
        getScore() -> int: Return risk score'''
doc.add_paragraph(spec)

# ============================================
# APPENDIX B: DETECTION RULES CATALOG
# ============================================
doc.add_page_break()
doc.add_heading('Appendix B: Detection Rules Catalog', 1)

doc.add_paragraph(
    'The following detection rules will be implemented in the initial release:'
)

rules_table = doc.add_table(rows=12, cols=4)
set_table_borders(rules_table)
rules_table.style = 'Table Grid'
rules_table.cell(0, 0).text = 'Rule ID'
rules_table.cell(0, 1).text = 'Rule Name'
rules_table.cell(0, 2).text = 'Detection Pattern'
rules_table.cell(0, 3).text = 'Severity'
for cell in rules_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

rules_data = [
    ('RULE-001', 'Brute Force Login', '5+ failed logins from same IP in 5 min', '8'),
    ('RULE-002', 'Port Scan Detection', '10+ connection attempts to different ports', '7'),
    ('RULE-003', 'Privilege Escalation', 'User gains admin/root privileges', '9'),
    ('RULE-004', 'After-Hours Access', 'Login outside business hours (6PM-6AM)', '4'),
    ('RULE-005', 'Multiple Failed SSH', '3+ failed SSH attempts in 1 min', '6'),
    ('RULE-006', 'Suspicious File Access', 'Access to sensitive paths (/etc/shadow)', '8'),
    ('RULE-007', 'Account Lockout', 'Account locked due to failed attempts', '5'),
    ('RULE-008', 'New Admin Account', 'New account created with admin rights', '9'),
    ('RULE-009', 'Service Stopped', 'Critical service stopped unexpectedly', '7'),
    ('RULE-010', 'Unusual Data Transfer', 'Large outbound data transfer detected', '6'),
    ('RULE-011', 'Multiple Source IPs', 'Same user logged in from 3+ IPs', '5'),
]

for i, row_data in enumerate(rules_data, 1):
    for j, cell_data in enumerate(row_data):
        rules_table.cell(i, j).text = cell_data

doc.add_paragraph()
cap = doc.add_paragraph('Table B.1: Detection Rules Catalog')
cap.runs[0].italic = True

# Save the document
doc.save(r'E:\CyberRisk\docs\CyberRisk_Monitor_Project_Design_Document.docx')
print('Document created successfully!')
print('Saved to: E:\\CyberRisk\\docs\\CyberRisk_Monitor_Project_Design_Document.docx')
